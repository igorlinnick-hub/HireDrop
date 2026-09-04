"""Generate ATS-compliant PDF resumes from raw resume text.

Format spec (per the ATS guide in the project):
  - Font: Helvetica (PDF-safe Arial equivalent)
  - Name: 16pt bold, ALL CAPS, left-aligned
  - Section headers: 12pt bold, left-aligned, with horizontal rule
  - Job title lines: 11pt bold, left-aligned
  - Body / bullets: 10.5pt, line spacing 1.0
  - Margins: 0.75 inch
  - Section spacing: 6pt between sections
  - Bullets: hyphen + space (no special symbols)
  - No tables, no images, no columns, no text boxes

Flow:
  1. Claude structures raw resume text into JSON
  2. reportlab renders structured JSON as a clean PDF
"""

import io
import json
import re

import pdfplumber
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from config import ANTHROPIC_API_KEY

SONNET_MODEL = "claude-sonnet-4-6"

_MARGIN = 0.75 * inch
_PAGE_WIDTH, _PAGE_HEIGHT = letter


def _docx_pt(points: float):
    from docx.shared import Pt

    return Pt(points)


def _docx_inch(inches: float):
    from docx.shared import Inches

    return Inches(inches)


def _extract_text_from_bytes(pdf_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        parts = []
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
    return "\n".join(parts)


ATS_PASS_THRESHOLD = 80


def clean_linkedin_url(url: str) -> str:
    """Normalize a LinkedIn URL to the clean, ATS-friendly form.

    LinkedIn appends an auto-generated suffix (digits/hash) to custom URLs when
    the vanity slug is taken, e.g. `linkedin.com/in/john-maxwell-9a8b7c6d5`.
    Recruiters and ATS prefer the bare `linkedin.com/in/john-maxwell`. We:
      - strip protocol (https://) and `www.`
      - drop a trailing `-<segment>` from the /in/ slug if that segment
        contains a digit (the tell-tale of an auto-suffix, not a real name part)
    A clean slug like `igor-lynnik` is left untouched (no digits → kept).
    """
    if not url:
        return url
    u = url.strip()
    u = re.sub(r"^https?://", "", u, flags=re.IGNORECASE)
    u = re.sub(r"^www\.", "", u, flags=re.IGNORECASE)
    u = u.rstrip("/")

    m = re.search(r"(linkedin\.com/in/)(.+)$", u, flags=re.IGNORECASE)
    if not m:
        return u
    prefix, slug = m.group(1), m.group(2)
    # Drop one or more trailing hash-like segments (segments containing a digit).
    parts = slug.split("-")
    while len(parts) > 1 and re.search(r"\d", parts[-1]):
        parts.pop()
    return prefix + "-".join(parts)


def get_ats_questions(resume_text: str) -> list[str]:
    """Ask Claude what ATS-critical keywords are implied but missing from this resume.

    ATS systems are keyword-matching engines — they scan for specific tool names,
    software, certifications, and job titles. This function generates 2-4 questions
    that surface implied skills and tools that should be explicitly listed so ATS
    can find them.

    Returns list of 2-4 questions. Empty list if nothing meaningful to ask.
    """
    if not ANTHROPIC_API_KEY or not resume_text.strip():
        return []

    try:
        import anthropic

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        prompt = f"""You are an ATS (Applicant Tracking System) optimization specialist.

ATS software is a keyword-matching engine — it scans resumes for EXACT words like tool names, software, certifications, and job titles. It does NOT care about achievements, soft skills, or how many people someone managed.

Review this resume and identify 2-4 questions that will surface ATS-critical keywords that are IMPLIED by the experience but NOT explicitly listed.

Focus ONLY on:
1. Specific tools/software/platforms implied by the work but not named (e.g., if they ran Instagram DM automation → ask which tool: ManyChat? Chatfuel? Manychat?)
2. Specific platforms used for paid ads (Meta Ads Manager? Google Ads? TikTok Ads?)
3. CRM or booking software (HubSpot? Salesforce? Mindbody? Calendly?)
4. Certifications with missing details (year, issuer, full name)
5. LinkedIn URL if not in contact section
6. Any specific industry software not listed (video editing: Premiere Pro? CapCut? DaVinci? | Design: Canva? Adobe Suite? Figma?)

DO NOT ask about:
- How many people they managed
- Revenue or budget numbers
- Soft skills
- Anything already clearly listed in the resume

Maximum 4 questions. Each question should be direct and specific — name the implied tool/category.

Resume:
{resume_text[:3000]}

Return ONLY a JSON array of question strings, no explanation:
["Question 1?", "Question 2?"]"""

        message = client.messages.create(
            model=SONNET_MODEL,
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        questions = [q for q in json.loads(raw) if isinstance(q, str)][:4]

        # LinkedIn URL is ATS-critical for recruiter verification. Always ask for it
        # if it's not already in the resume text (common when the original was a
        # design PDF with LinkedIn in a graphical column pdfplumber can't read).
        has_linkedin_url = bool(re.search(r"linkedin\.com/in/", resume_text, re.IGNORECASE))
        if not has_linkedin_url:
            linkedin_q = "What is your LinkedIn URL? (e.g., linkedin.com/in/your-name)"
            if not any("linkedin" in q.lower() for q in questions):
                # Prepend it; cap total at 4 by dropping the last AI question if needed
                questions = [linkedin_q] + questions[:3]

        return questions
    except Exception as e:
        print(f"[ats_questions] Failed: {e}")
        return []


def _structure_resume(resume_text: str, answers: list[dict] | None = None) -> dict:
    """Ask Claude to parse resume text into a structured JSON object."""
    if not ANTHROPIC_API_KEY:
        return {}

    try:
        import anthropic

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        answers_block = ""
        if answers:
            qa_lines = "\n".join(
                f"Q: {a.get('question', '')}\nA: {a.get('answer', '')}"
                for a in answers
                if a.get("answer", "").strip()
            )
            if qa_lines:
                answers_block = f"""

ADDITIONAL ATS KEYWORDS FROM CANDIDATE (these are specific tools, software, platforms, and certifications the candidate confirmed they used):
{qa_lines}

IMPORTANT: Extract every specific tool/software/platform name from the answers above and add them to the "tech_skills" array. Also add them to "competencies" if relevant. Do NOT leave implied tools unnamed — if the answer names a specific tool, it MUST appear verbatim in tech_skills."""

        prompt = f"""Extract the resume below into this exact JSON structure.
Return ONLY the JSON object, no explanation, no markdown.
Incorporate any additional candidate answers into the appropriate fields (add metrics to bullets, fill contact, etc.).

{{
  "name": "FULL NAME",
  "title": "Professional title / tagline (if present, else empty string)",
  "contact": {{
    "phone": "",
    "email": "",
    "location": "",
    "linkedin": ""
  }},
  "summary": "Professional summary paragraph (plain text, no bullets)",
  "competencies": ["Competency 1", "Competency 2"],
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, State",
      "dates": "Month Year – Month Year",
      "bullets": ["Bullet point 1", "Bullet point 2"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree Name",
      "school": "School Name",
      "year": "Year (or empty)"
    }}
  ],
  "certifications": ["Certification 1"],
  "languages": ["Language (Level)", "Language (Level)"],
  "tech_skills": ["Skill 1", "Skill 2"]
}}

RESUME TEXT:
{resume_text[:4000]}{answers_block}"""

        message = client.messages.create(
            model=SONNET_MODEL,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        data = json.loads(raw)
        # ATS hygiene: normalize the LinkedIn URL (strip www/protocol + hash suffix)
        contact = data.get("contact") or {}
        if contact.get("linkedin"):
            contact["linkedin"] = clean_linkedin_url(contact["linkedin"])
            data["contact"] = contact
        return data
    except Exception as e:
        print(f"[ats_pdf] Structure extraction failed: {e}")
        return {}


def structure_resume_data(resume_text: str, answers: list[dict] | None = None) -> dict:
    """Public wrapper: structure resume text into the ATS JSON, once.

    Both the PDF and DOCX renderers consume this so we make a single Claude call
    when generating both formats. Falls back to a first-line name if extraction
    returns nothing usable.
    """
    data = _structure_resume(resume_text, answers=answers or [])
    if not data.get("name"):
        lines = [ln.strip() for ln in (resume_text or "").split("\n") if ln.strip()]
        data["name"] = lines[0] if lines else "CANDIDATE"
    return data


def _make_styles() -> dict:
    base = ParagraphStyle(
        "base",
        fontName="Helvetica",
        fontSize=10.5,
        leading=12.6,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=0,
    )
    return {
        "name": ParagraphStyle(
            "name",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            spaceAfter=2,
            textTransform="uppercase",
        ),
        "title_line": ParagraphStyle(
            "title_line",
            parent=base,
            fontSize=10.5,
            leading=13,
            spaceAfter=1,
        ),
        "contact": ParagraphStyle(
            "contact",
            parent=base,
            fontSize=10.5,
            leading=13,
            spaceAfter=0,
        ),
        "section_header": ParagraphStyle(
            "section_header",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            spaceBefore=6,
            spaceAfter=2,
        ),
        "job_title": ParagraphStyle(
            "job_title",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=13,
            spaceBefore=3,
            spaceAfter=0,
        ),
        "job_meta": ParagraphStyle(
            "job_meta",
            parent=base,
            fontSize=10.5,
            leading=12.6,
            spaceAfter=1,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=base,
            fontSize=10.5,
            leading=12.6,
            leftIndent=12,
            spaceAfter=0,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base,
            fontSize=10.5,
            leading=12.6,
            spaceAfter=2,
        ),
        "competency_row": ParagraphStyle(
            "competency_row",
            parent=base,
            fontSize=10.5,
            leading=13,
            spaceAfter=0,
        ),
    }


def _section_block(title: str, styles: dict) -> list:
    return [
        Spacer(1, 4),
        Paragraph(title.upper(), styles["section_header"]),
        HRFlowable(width="100%", thickness=0.75, color=colors.black, spaceAfter=3),
    ]


def _build_story(data: dict, styles: dict) -> list:
    story = []

    # Name
    name = (data.get("name") or "").upper()
    story.append(Paragraph(name, styles["name"]))

    # Optional title line
    if data.get("title"):
        story.append(Paragraph(data["title"], styles["title_line"]))

    # Contact line
    c = data.get("contact") or {}
    contact_parts = [
        p for p in [c.get("phone"), c.get("email"), c.get("location"), c.get("linkedin")] if p
    ]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), styles["contact"]))

    story.append(Spacer(1, 4))

    # Professional Summary
    if data.get("summary"):
        story.extend(_section_block("Professional Summary", styles))
        story.append(Paragraph(data["summary"], styles["body"]))

    # Core Competencies
    comps = data.get("competencies") or []
    if comps:
        story.extend(_section_block("Core Competencies", styles))
        # Render as pipe-separated rows, max ~4 per row
        chunk_size = 4
        chunks = [comps[i : i + chunk_size] for i in range(0, len(comps), chunk_size)]
        for chunk in chunks:
            story.append(Paragraph(" | ".join(chunk), styles["competency_row"]))

    # Professional Experience
    exp = data.get("experience") or []
    if exp:
        story.extend(_section_block("Professional Experience", styles))
        for job in exp:
            job_title = job.get("title") or ""
            company = job.get("company") or ""
            location = job.get("location") or ""
            dates = job.get("dates") or ""

            meta_parts = [p for p in [company, location, dates] if p]
            header_text = f"<b>{job_title}</b>"
            if meta_parts:
                header_text += f"  |  {' — '.join(meta_parts)}"
            story.append(Paragraph(header_text, styles["job_title"]))

            for bullet in job.get("bullets") or []:
                bullet_text = bullet.lstrip("-•– ").strip()
                story.append(Paragraph(f"- {bullet_text}", styles["bullet"]))
            story.append(Spacer(1, 3))

    # Education & Certifications
    edu = data.get("education") or []
    certs = data.get("certifications") or []
    if edu or certs:
        story.extend(_section_block("Education & Certifications", styles))
        for e in edu:
            degree = e.get("degree") or ""
            school = e.get("school") or ""
            year = e.get("year") or ""
            parts = [p for p in [school, year] if p]
            line = f"<b>{degree}</b>"
            if parts:
                line += f" — {' | '.join(parts)}"
            story.append(Paragraph(line, styles["body"]))
        for cert in certs:
            story.append(Paragraph(f"<b>{cert}</b>", styles["body"]))

    # Tech Skills
    tech = data.get("tech_skills") or []
    if tech:
        story.extend(_section_block("Technical Skills", styles))
        chunk_size = 4
        chunks = [tech[i : i + chunk_size] for i in range(0, len(tech), chunk_size)]
        for chunk in chunks:
            story.append(Paragraph(" | ".join(chunk), styles["competency_row"]))

    # Languages
    langs = data.get("languages") or []
    if langs:
        story.extend(_section_block("Languages", styles))
        story.append(Paragraph(" | ".join(langs), styles["body"]))

    return story


def _resolve_data(
    pdf_bytes: bytes | None,
    resume_text: str | None,
    answers: list[dict] | None,
    data: dict | None,
) -> dict:
    """Return structured resume data, building it once if not supplied."""
    if data is not None:
        return data
    if pdf_bytes and not resume_text:
        resume_text = _extract_text_from_bytes(pdf_bytes)
    if not resume_text or not resume_text.strip():
        raise ValueError("No resume text to process")
    return structure_resume_data(resume_text, answers=answers or [])


def generate_ats_pdf(
    pdf_bytes: bytes | None = None,
    resume_text: str | None = None,
    answers: list[dict] | None = None,
    data: dict | None = None,
) -> bytes:
    """Generate an ATS-compliant PDF.

    Args:
        pdf_bytes: original PDF (text extracted with pdfplumber)
        resume_text: pre-extracted plain text (alternative to pdf_bytes)
        answers: list of {question, answer} dicts from user Q&A — merged into the resume
        data: pre-structured resume dict (skips the Claude call — use when also
              rendering DOCX from the same structuring)

    Returns PDF bytes.
    """
    data = _resolve_data(pdf_bytes, resume_text, answers, data)

    styles = _make_styles()
    story = _build_story(data, styles)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=_MARGIN,
        rightMargin=_MARGIN,
        topMargin=_MARGIN,
        bottomMargin=_MARGIN,
        title=f"{data.get('name', 'Resume')} — ATS Resume",
    )
    doc.build(story)
    return buffer.getvalue()


def _docx_section_header(doc, title: str) -> None:
    """Add a 12pt bold uppercase section header with a bottom rule."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = _docx_pt(8)
    p.paragraph_format.space_after = _docx_pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = _docx_pt(12)
    run.font.name = "Arial"
    # Bottom border (horizontal rule)
    pPr = p._p.get_or_add_pPr()  # noqa: N806 — OOXML element name, keep the spec spelling
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)


def generate_ats_docx(
    pdf_bytes: bytes | None = None,
    resume_text: str | None = None,
    answers: list[dict] | None = None,
    data: dict | None = None,
) -> bytes:
    """Generate an ATS-compliant .docx mirroring the PDF layout.

    Some job boards / employers only accept Word documents, so we offer DOCX
    alongside PDF. Same structured data, same section order, ATS-safe formatting
    (Arial, hyphen bullets, no tables/columns/images).
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data = _resolve_data(pdf_bytes, resume_text, answers, data)

    doc = Document()
    # Base font + 0.75" margins
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = _docx_inch(0.75)
        section.left_margin = section.right_margin = _docx_inch(0.75)

    def add_line(text, *, size=10.5, bold=False, space_after=0, space_before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Arial"
        return p

    # Name + title + contact
    add_line((data.get("name") or "").upper(), size=16, bold=True, space_after=2)
    if data.get("title"):
        add_line(data["title"], size=10.5, space_after=1)
    c = data.get("contact") or {}
    contact_parts = [
        p for p in [c.get("phone"), c.get("email"), c.get("location"), c.get("linkedin")] if p
    ]
    if contact_parts:
        add_line(" | ".join(contact_parts), size=10.5)

    # Professional Summary
    if data.get("summary"):
        _docx_section_header(doc, "Professional Summary")
        add_line(data["summary"], space_after=2)

    # Core Competencies
    comps = data.get("competencies") or []
    if comps:
        _docx_section_header(doc, "Core Competencies")
        for i in range(0, len(comps), 4):
            add_line(" | ".join(comps[i : i + 4]))

    # Professional Experience
    exp = data.get("experience") or []
    if exp:
        _docx_section_header(doc, "Professional Experience")
        for job in exp:
            meta = [p for p in [job.get("company"), job.get("location"), job.get("dates")] if p]
            header = job.get("title") or ""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(header)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Arial"
            if meta:
                r2 = p.add_run("  |  " + " — ".join(meta))
                r2.font.size = Pt(10.5)
                r2.font.name = "Arial"
            for bullet in job.get("bullets") or []:
                bt = bullet.lstrip("-•– ").strip()
                bp = add_line(f"- {bt}")
                bp.paragraph_format.left_indent = _docx_pt(12)

    # Education & Certifications
    edu = data.get("education") or []
    certs = data.get("certifications") or []
    if edu or certs:
        _docx_section_header(doc, "Education & Certifications")
        for e in edu:
            parts = [p for p in [e.get("school"), e.get("year")] if p]
            line = e.get("degree") or ""
            if parts:
                line += " — " + " | ".join(parts)
            add_line(line, bold=True)
        for cert in certs:
            add_line(cert, bold=True)

    # Technical Skills
    tech = data.get("tech_skills") or []
    if tech:
        _docx_section_header(doc, "Technical Skills")
        for i in range(0, len(tech), 4):
            add_line(" | ".join(tech[i : i + 4]))

    # Languages
    langs = data.get("languages") or []
    if langs:
        _docx_section_header(doc, "Languages")
        add_line(" | ".join(langs))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
